from docx import Document
from docx_extractor import extract_docx_text


def test_docx_extraction(tmp_path):
    file_path = tmp_path / "test_resume.docx"

    document = Document()
    document.add_paragraph("Rounak Kumar")
    document.add_paragraph("Python FastAPI Machine Learning")
    document.save(file_path)

    text = extract_docx_text(str(file_path))

    assert "Rounak Kumar" in text
    assert "Python" in text
    assert "FastAPI" in text
    assert "Machine Learning" in text